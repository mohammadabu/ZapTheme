# -*- coding: utf-8 -*-

from sys import platform
from subprocess import Popen
from docx2pdf import convert
from datetime import datetime, date
import os
import base64
import logging
from odoo import models, tools
from odoo.addons.sale_export_docx.reports import template
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Inches , Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import docx
_logger = logging.getLogger(__name__)

try:
    from docxtpl import DocxTemplate
except ImportError:
    _logger.debug('Can not import DocxTemplate')


class ExportDocxAbstract(models.AbstractModel):
    _name = 'export.docx.abstract'
    _description = 'Export Docx Abstract Model'

    def _get_objs_for_report(self, docids, data):
        if docids:
            ids = docids
        elif data and 'context' in data:
            ids = data["context"].get('active_ids', [])
        else:
            ids = self.env.context.get('active_ids', [])
        return self.env[self.env.context.get('active_model')].browse(ids)

    def create_docx_report(self, docids, data):
        objs = self._get_objs_for_report(docids, data)
        return self.generate_docx_report( data, objs), 'docx'

    def generate_docx_report(self, data, objs):
        timestamp = str(int(datetime.timestamp(datetime.now())))
        path_docx = '/var/lib/odoo/.local/share/Odoo/'






        document = docx.Document()
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(3.3)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        menuTable = document.add_table(rows=1,cols=3)
        hdr_cells = menuTable.rows[0].cells

        hdr_cells[0].width = Inches(2)
        hdr_cells[0].text = 'Salary Certificate'
        paragraph = hdr_cells[0].paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size= Pt(16)
        font.bold = True
        font.name = 'Calibri (Body)'
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



        hdr_cells[1].width = Inches(0.1)


        hdr_cells[2].width = Inches(3)
        hdr_cells[2].text = 'شهادة تعريف بالراتب'
        paragraph_1 = hdr_cells[2].paragraphs[0]
        run_1 = paragraph_1.runs
        font_1 = run_1[0].font
        font_1.size= Pt(16)
        font_1.name = 'Al-Mateen'
        paragraph_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # par_0 = hdr_Cells[0].paragraphs[0] = 'Salary Certificate'
        # run_0 = par_0.runs
        # font_0 = run_0[0].font
        # font_0.size= Pt(30)

        # font_paragraph_0 = paragraph_0.style.font
        # font_paragraph_0.name = 'Times New Roman'
        # font_paragraph_0.size = Pt(16)
        # font_paragraph_0.bold = True
        # font_0 = run_0.font
        # font_0.name = 'Times New Roman'
        # font_0.size = Pt(16)


        # paragraph_1 = hdr_Cells[1].add_paragraph('شهادة تعريف بالراتب')
        # font_paragraph_1 = paragraph_1.style.font
        # font_paragraph_1.name = 'Calibri'
        # font_paragraph_1.size = Pt(12)
        # font_paragraph_1.bold = True
        # run = paragraph_0.add_run('Salary Certificate')
        # run.bold = True
        # run.size = Pt(16)
        # run_font = run.style.font
        # run_font.name = 'Times New Roman'
        # paragraph_1 = hdr_Cells[1].add_paragraph('شهادة تعريف بالراتب')
        # paragraph.add_run(' sit amet.')
        # hdr_Cells_text_0 = hdr_Cells[0].text = 'Salary Certificate'
        # paragraph_0 = hdr_Cells_text_0.paragraphs[0]
        # run_0 = paragraph_0.runs
        # font_0 = run_0[0].font
        # font_0.size= Pt(30)
        # font_hdr_0 = hdr_Cells_0_parag.style.font
        # font_hdr_0.name = 'Times New Roman'
        # font_hdr_0.size = Pt(16)
        # font_hdr_0.bold = True





        
        # hdr_Cells[1].add_paragraph('شهادة تعريف بالراتب')
        # records = [
        #     ['To :','الى : ']
        # ]
        # for ID, nameOfMeal in records:
        #     row_Cells = menuTable.add_row().cells
        #     row_Cells[0].add_paragraph(str(ID))
        #     row_Cells[1].add_paragraph(nameOfMeal)



        path_docx = path_docx + '/EmployeeDocx_' + timestamp + ".docx"
        document.save(path_docx)


        
        report_doxc_path = path_docx
        with open(report_doxc_path, mode='rb') as file:
            fileContent = file.read()

        try:
            os.remove(report_doxc_path)
        except Exception as e:
            _logger.warning(repr(e))

        return fileContent

    def generate_variables(self, objs):
        raise NotImplementedError()

    def get_report_name(self, objs):
        raise NotImplementedError()

    def _save_file(self, template_path, data):
        out_stream = open(template_path, 'wb')
        try:
            out_stream.write(data)
        finally:
            out_stream.close()

    def get_partner_address(self, obj=None):
        if not obj:
            return ''
        address = ''
        address += f'{obj.street}' if obj.street else ''
        address += f', {obj.street}' if obj.street2 else ''
        address += f', {obj.city}' if obj.city else ''
        address += f', {obj.state_id.name}' if obj.state_id else ''
        address += f', {obj.country_id.name}' if obj.country_id else ''
        return address
